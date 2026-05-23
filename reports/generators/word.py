"""
Word document report generator using python-docx.
"""
import io

from docx import Document
from docx.shared import Pt, RGBColor


def build_summary_docx(title: str, rows: list[tuple], narrative: str = '') -> bytes:
    """
    Build a simple Word document report.

    Args:
        title: Document title.
        rows: List of (label, value) tuples.
        narrative: Optional narrative text.

    Returns:
        DOCX file content as bytes.
    """
    doc = Document()

    heading = doc.add_heading(title, level=1)
    heading.runs[0].font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)

    if rows:
        doc.add_heading('Summary Table', level=2)
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text = 'Indicator'
        hdr[1].text = 'Value'
        for label, value in rows:
            row = table.add_row().cells
            row[0].text = str(label)
            row[1].text = str(value)

    if narrative:
        doc.add_heading('Narrative Summary', level=2)
        for para in narrative.split('\n\n'):
            if para.strip():
                doc.add_paragraph(para.strip())

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
