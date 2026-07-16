import sys
from docx import Document
sys.stdout.reconfigure(encoding='utf-8')
doc = Document(r'C:\Users\HP\Downloads\January to June Action Plan_2026_Kurigram.docx')
t1, t2 = doc.tables[0], doc.tables[1]
print('T1 R14 c0:', repr(t1.rows[14].cells[0].text))
print('T1 R18 c0:', repr(t1.rows[18].cells[0].text))
print('T1 R10 c0:', repr(t1.rows[10].cells[0].text))
print('T2 R00 c0:', repr(t2.rows[0].cells[0].text))
print('T2 R07 c0:', repr(t2.rows[7].cells[0].text))
print('T2 R00 c1:', repr(t2.rows[0].cells[1].text))
print('T1 R01 c1:', repr(t1.rows[1].cells[1].text))
