import sys, openpyxl
from pathlib import Path
sys.stdout.reconfigure(encoding='utf-8')

SRC = r'C:\Users\HP\Downloads\MIS Tools_100626.xlsx'
OUT = r'C:\Users\HP\Documents\spondon_clone\_bandhu_source_dump.md'

wb = openpyxl.load_workbook(SRC, data_only=True)
print('SHEETS:', wb.sheetnames)
lines = [f'# MIS Tools_100626.xlsx — {len(wb.sheetnames)} sheets\n']
for ws in wb.worksheets:
    lines.append(f'\n\n## SHEET: {ws.title}  ({ws.max_row}x{ws.max_column})\n')
    for r in range(1, ws.max_row + 1):
        cells = []
        for c in range(1, ws.max_column + 1):
            v = ws.cell(r, c).value
            if v is not None and str(v).strip():
                cells.append(str(v).strip())
        if cells:
            lines.append(' | '.join(cells))
    print(f'  {ws.title}: {ws.max_row}x{ws.max_column}')
Path(OUT).write_text('\n'.join(lines), encoding='utf-8')
print('WROTE', OUT, len('\n'.join(lines)), 'chars')

try:
    from docx import Document
    for src, out in [
        (r'C:\Users\HP\Downloads\M&E Tools correction.docx',
         r'C:\Users\HP\Documents\spondon_clone\_bandhu_correction1.md'),
        (r'C:\Users\HP\Downloads\Correction for SIMPLE_AS.docx',
         r'C:\Users\HP\Documents\spondon_clone\_bandhu_correction2.md'),
    ]:
        d = Document(src)
        txt = []
        for p in d.paragraphs:
            if p.text.strip():
                txt.append(p.text)
        for ti, t in enumerate(d.tables):
            txt.append(f'\n[TABLE {ti}]')
            for row in t.rows:
                txt.append(' | '.join(c.text.strip() for c in row.cells))
        Path(out).write_text('\n'.join(txt), encoding='utf-8')
        print('WROTE', out, len('\n'.join(txt)), 'chars')
except Exception as e:
    print('docx extract failed:', repr(e))
