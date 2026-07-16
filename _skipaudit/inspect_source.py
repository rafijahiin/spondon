import sys, re, os
sys.stdout.reconfigure(encoding='utf-8')

# 1) LGBTQI (Hijra) source docx — how are skips / specify expressed?
try:
    from docx import Document
    d = Document(r'C:\Users\HP\Downloads\Baseline_Questionnaire_LGBTQI_Population_NK_formatted_NK.docx')
    text = '\n'.join(p.text for p in d.paragraphs)
    for t in d.tables:
        for row in t.rows:
            text += '\n' + ' | '.join(c.text for c in row.cells)
    print('=== LGBTQI docx: skip-instruction lines ===')
    pat = re.compile(r'(skip|go to|→|proceed to|if .*(then|→|go)|otherwise)', re.I)
    hits = [ln.strip() for ln in text.splitlines() if pat.search(ln) and len(ln.strip()) < 220]
    for h in hits[:25]:
        print('  •', h[:200])
    print(f'  ...total skip-like lines: {len(hits)}')
    spec = [ln.strip() for ln in text.splitlines() if re.search(r'\(specify\)|other \(specify\)', ln, re.I)]
    print(f'  "(specify)" mentions: {len(spec)}   e.g. {spec[:3]}')
except Exception as e:
    print('docx read failed:', e)

# 2) FSW data-collection-plan xlsx — does it carry structured skip logic?
try:
    import openpyxl
    wb = openpyxl.load_workbook(r'C:\Users\HP\Downloads\Data_Collection_Plans_FSW_and_TGD-1.xlsx', read_only=True)
    print('\n=== Data_Collection_Plans_FSW xlsx ===')
    print('  sheets:', wb.sheetnames)
    for sh in wb.sheetnames[:3]:
        ws = wb[sh]
        hdr = [c.value for c in next(ws.iter_rows(max_row=1))]
        print(f'  [{sh}] columns:', [str(h)[:30] for h in hdr if h])
except Exception as e:
    print('xlsx read failed:', e)
