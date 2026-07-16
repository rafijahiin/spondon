import sys, re
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
p = r'C:\Users\HP\Downloads\Baseline_Questionnaire_LGBTQI_Population_NK_formatted_NK.docx'
doc = Document(p)
# Gather all text from paragraphs AND tables
texts = []
for para in doc.paragraphs:
    if para.text.strip():
        texts.append(para.text.strip())
for tbl in doc.tables:
    for row in tbl.rows:
        for cell in row.cells:
            t = cell.text.strip()
            if t:
                texts.append(t)
blob = '\n'.join(texts)
print('total text blocks:', len(texts), '| chars:', len(blob))
# Count skip-logic and specify markers
def cnt(pat): return len(re.findall(pat, blob, re.I))
print('\n--- markers of explicit logic in the SOURCE ---')
print('  "skip" occurrences        :', cnt(r'\bskip\b'))
print('  "go to" / "goto"          :', cnt(r'go\s*to'))
print('  arrow "→"                 :', blob.count('→'))
print('  "(specify)" / "specify"   :', cnt(r'specify'))
print('  "if ... " conditionals    :', cnt(r'\bif\b'))
print('  "not applicable"/"N/A"    :', cnt(r'not applicable|n/a'))
# Show sample lines that carry skip / specify instructions
print('\n--- sample SOURCE lines with skip/specify/arrow ---')
shown = 0
for t in texts:
    if re.search(r'\bskip\b|go\s*to|→|\(specify\)', t, re.I) and 8 < len(t) < 240:
        print('  •', t)
        shown += 1
        if shown >= 25: break
